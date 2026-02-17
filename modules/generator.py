import os
import subprocess
import jinja2
from datetime import date
from docx import Document as DocxDocument
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

def escape_latex(text):
    """Escapes strings for LaTeX safety."""
    if not isinstance(text, str):
        return text
    chars = {
        '&': r'\&', '%': r'\%', '$': r'\$', '#': r'\#', '_': r'\_',
        '{': r'\{', '}': r'\}', '~': r'\textasciitilde{}', '^': r'\^{}',
        '\\': r'\textbackslash{}', '<': r'\textless{}', '>': r'\textgreater{}',
        "'": r"'", 
    }
    return "".join(chars.get(c, c) for c in text)

def generate_resume_pdf(data, template_name="modern", output_dir="output"):
    """Generates PDF from LaTeX template."""
    
    # Setup Jinja2 for LaTeX
    template_loader = jinja2.FileSystemLoader(searchpath="./assets/templates")
    latex_jinja_env = jinja2.Environment(
        loader=template_loader,
        block_start_string='\\BLOCK{', block_end_string='}',
        variable_start_string='\\VAR{', variable_end_string='}',
        comment_start_string='\\#{', comment_end_string='}',
        line_statement_prefix='%%', line_comment_prefix='%#',
        trim_blocks=True, autoescape=False,
    )
    
    try:
        template = latex_jinja_env.get_template(f"{template_name}.tex")
    except jinja2.TemplateNotFound:
        template = latex_jinja_env.get_template('modern.tex')

    # Data Cleaning & Preparation
    clean_data = {}
    
    def clean_structure(item):
        if isinstance(item, list):
            return [clean_structure(i) for i in item]
        elif isinstance(item, dict):
            return {k: clean_structure(v) for k, v in item.items()}
        elif isinstance(item, str):
            return escape_latex(item)
        return item

    clean_data = clean_structure(data)
    
    # Ensure skills format
    if 'skills' in clean_data:
        # Handle various formats AI might have returned
        cleaned_skills = []
        for s in clean_data['skills']:
            if isinstance(s.get('items'), list):
                s['items'] = ', '.join(s['items'])
            cleaned_skills.append(s)
        clean_data['skills'] = cleaned_skills

    # Render LaTeX
    rendered_tex = template.render(**clean_data, today=date.today().strftime("%B %Y"))

    if not os.path.exists(output_dir):
        os.makedirs(output_dir)
        
    tex_path = os.path.join(output_dir, "resume.tex")
    with open(tex_path, "w", encoding='utf-8') as f:
        f.write(rendered_tex)

    # Compile LaTeX
    try:
        # Run pdflatex twice for layout (usually needed) but once is fine for simple templates
        result = subprocess.run(
            ['pdflatex', '-interaction=nonstopmode', f'-output-directory={output_dir}', tex_path],
            capture_output=True,
            text=True,
            timeout=30
        )
        
        if result.returncode != 0:
            error_msg = f"LaTeX compilation failed (Code {result.returncode}). Check if 'pdflatex' is installed and valid."
            # Only show first few lines of error to avoid overwhelming user
            if result.stdout:
                error_msg += f"\nLast log: {result.stdout[-200:]}"
            raise Exception(error_msg)
            
        pdf_path = os.path.join(output_dir, "resume.pdf")
        return pdf_path
        
    except FileNotFoundError:
        raise Exception("System Error: 'pdflatex' not found. Please install a LaTeX distribution (TeX Live/MiKTeX).")
    except subprocess.TimeoutExpired:
        raise Exception("Timeout: PDF generation took too long.")

def add_bottom_border(paragraph):
    """Helper to add bottom border to Word headings."""
    p = paragraph._p
    pPr = p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '6')
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), '000000')
    pBdr.append(bottom)
    pPr.append(pBdr)

def generate_resume_docx(data, output_dir="output"):
    """Generates a professional Word document."""
    doc = DocxDocument()
    
    # Setup Page
    section = doc.sections[0]
    section.top_margin = Inches(0.5)
    section.bottom_margin = Inches(0.5)
    section.left_margin = Inches(0.5)
    section.right_margin = Inches(0.5)
    
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Times New Roman'
    font.size = Pt(11)
    
    # Header
    name_para = doc.add_paragraph()
    name_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    name_run = name_para.add_run(data.get('name', 'Name'))
    name_run.bold = True
    name_run.font.size = Pt(24)
    name_run.font.name = 'Times New Roman'
    
    # Contact Info
    contact_parts = []
    for field in ['email', 'phone', 'linkedin', 'github', 'website']:
        if data.get(field):
            val = data.get(field)
            if field == 'linkedin' and 'linkedin' not in val.lower(): val = f"LinkedIn: {val}"
            if field == 'github' and 'github' not in val.lower(): val = f"GitHub: {val}"
            contact_parts.append(val)
    
    if contact_parts:
        contact_para = doc.add_paragraph(" | ".join(contact_parts))
        contact_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        contact_para.paragraph_format.space_after = Pt(10)

    def add_section(title):
        p = doc.add_heading(title.upper(), level=1)
        run = p.runs[0]
        run.font.name = 'Times New Roman'
        run.font.size = Pt(12)
        run.font.color.rgb = RGBColor(0, 0, 0)
        add_bottom_border(p)
        p.paragraph_format.space_before = Pt(12)
        p.paragraph_format.space_after = Pt(6)

    # Summary
    if data.get('summary'):
        add_section('Professional Summary')
        doc.add_paragraph(data.get('summary'))

    # Experience
    if data.get('experience'):
        add_section('Experience')
        for job in data.get('experience', []):
            table = doc.add_table(rows=1, cols=2)
            table.autofit = False
            
            c1 = table.cell(0, 0)
            c1.width = Inches(5.5)
            p1 = c1.paragraphs[0]
            p1.add_run(f"{job.get('title', 'Role')}").bold = True
            p1.add_run(f" | {job.get('company', 'Company')}").italic = True
            
            c2 = table.cell(0, 1)
            c2.width = Inches(2.0)
            p2 = c2.paragraphs[0]
            p2.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            p2.add_run(job.get('dates', '')).italic = True
            
            for bullet in job.get('bullets', []):
                doc.add_paragraph(bullet, style='List Bullet').paragraph_format.space_after = Pt(0)

    # Projects
    if data.get('projects'):
        add_section('Projects')
        for proj in data.get('projects', []):
            table = doc.add_table(rows=1, cols=2)
            table.autofit = False
            
            c1 = table.cell(0, 0)
            c1.width = Inches(5.5)
            p1 = c1.paragraphs[0]
            p1.add_run(proj.get('name', 'Project')).bold = True
            
            c2 = table.cell(0, 1)
            c2.width = Inches(2.0)
            p2 = c2.paragraphs[0]
            p2.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            if proj.get('link'):
                p2.add_run(proj.get('link')).italic = True

            if proj.get('description'):
                doc.add_paragraph(proj.get('description'), style='List Bullet').paragraph_format.space_after = Pt(2)

    # Education
    if data.get('education'):
        add_section('Education')
        for edu in data.get('education', []):
            table = doc.add_table(rows=1, cols=2)
            c1 = table.cell(0,0)
            c1.width = Inches(5.5)
            c1.paragraphs[0].add_run(edu.get('school', '')).bold = True
            
            c2 = table.cell(0,1)
            c2.width = Inches(2.0)
            c2.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
            c2.paragraphs[0].add_run(edu.get('year', ''))
            
            p = doc.add_paragraph()
            p.add_run(edu.get('degree', '')).italic = True
            if edu.get('gpa'):
                p.add_run(f" | GPA: {edu.get('gpa')}")
            p.paragraph_format.space_after = Pt(6)

    # Skills
    if data.get('skills'):
        add_section('Technical Skills')
        for skill in data.get('skills', []):
            p = doc.add_paragraph()
            p.paragraph_format.space_after = Pt(0)
            p.add_run(f"{skill.get('category')}: ").bold = True
            
            items = skill.get('items')
            if isinstance(items, list):
                items = ", ".join(items)
            p.add_run(str(items))

    if not os.path.exists(output_dir):
        os.makedirs(output_dir)
        
    docx_path = os.path.join(output_dir, "Optimized_Resume.docx")
    doc.save(docx_path)
    return docx_path
